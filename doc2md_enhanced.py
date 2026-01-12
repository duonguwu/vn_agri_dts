#!/usr/bin/env python3
"""
Enhanced Document to Markdown Converter
Hỗ trợ: PDF, DOC, DOCX -> Markdown
Author: Enhanced for Data Mining preprocessing
"""

import pymupdf4llm
import pathlib
import os
import argparse
import logging
from datetime import datetime
from typing import Optional, List, Dict
import sys
from concurrent.futures import ProcessPoolExecutor, as_completed
import multiprocessing

# Thêm thư viện cho DOC/DOCX
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  Warning: python-docx chưa được cài đặt. Chạy: pip install python-docx")

try:
    import pypandoc
    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False
    print("⚠️  Warning: pypandoc chưa được cài đặt. Chạy: pip install pypandoc")

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('conversion.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


class DocumentConverter:
    """Class để convert documents sang Markdown"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.doc', '.docx'}
    
    def __init__(self, output_dir: Optional[str] = None, preserve_structure: bool = True, 
                 skip_existing: bool = False):
        """
        Args:
            output_dir: Thư mục output (None = cùng thư mục với input)
            preserve_structure: Giữ nguyên cấu trúc thư mục
            skip_existing: Bỏ qua file đã convert
        """
        self.output_dir = output_dir
        self.preserve_structure = preserve_structure
        self.skip_existing = skip_existing
        self.stats = {
            'total': 0,
            'success': 0,
            'failed': 0,
            'skipped': 0
        }
    
    def convert_pdf_to_md(self, pdf_path: str, output_path: str) -> bool:
        """Convert PDF sang Markdown"""
        try:
            logger.info(f"📄 Converting PDF: {pdf_path}")
            md_text = pymupdf4llm.to_markdown(pdf_path)
            
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(md_text)
            
            logger.info(f"✅ Success: {output_path}")
            return True
        except Exception as e:
            logger.error(f"❌ Failed to convert {pdf_path}: {str(e)}")
            return False
    
    def convert_docx_to_md(self, docx_path: str, output_path: str) -> bool:
        """Convert DOCX sang Markdown (sử dụng python-docx)"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx không có sẵn")
            return False
        
        try:
            logger.info(f"📝 Converting DOCX: {docx_path}")
            doc = Document(docx_path)
            
            # Extract text với formatting cơ bản
            md_lines = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    md_lines.append("")
                    continue
                
                # Detect headings
                if para.style.name.startswith('Heading'):
                    level = para.style.name.split()[-1]
                    if level.isdigit():
                        md_lines.append(f"{'#' * int(level)} {text}")
                    else:
                        md_lines.append(f"## {text}")
                else:
                    md_lines.append(text)
            
            # Extract tables
            for table in doc.tables:
                md_lines.append("\n")
                for i, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    md_lines.append("| " + " | ".join(cells) + " |")
                    if i == 0:  # Header separator
                        md_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
                md_lines.append("\n")
            
            with open(output_path, "w", encoding="utf-8") as f:
                f.write("\n".join(md_lines))
            
            logger.info(f"✅ Success: {output_path}")
            return True
        except Exception as e:
            logger.error(f"❌ Failed to convert {docx_path}: {str(e)}")
            return False
    
    def convert_doc_to_md(self, doc_path: str, output_path: str) -> bool:
        """Convert DOC sang Markdown (sử dụng pypandoc)"""
        if not PANDOC_AVAILABLE:
            logger.warning("pypandoc không có sẵn, thử dùng python-docx...")
            # Fallback: try to open as docx
            return self.convert_docx_to_md(doc_path, output_path)
        
        try:
            logger.info(f"📃 Converting DOC: {doc_path}")
            # Sử dụng pandoc để convert
            output = pypandoc.convert_file(doc_path, 'md', outputfile=output_path)
            logger.info(f"✅ Success: {output_path}")
            return True
        except Exception as e:
            logger.error(f"❌ Failed to convert {doc_path}: {str(e)}")
            # Fallback
            logger.info("Trying fallback method...")
            return self.convert_docx_to_md(doc_path, output_path)
    
    def get_output_path(self, input_path: str, base_dir: Optional[str] = None) -> str:
        """Tính toán đường dẫn output"""
        input_path_obj = pathlib.Path(input_path)
        
        # Đổi extension sang .md
        md_filename = input_path_obj.stem + ".md"
        
        if self.output_dir is None:
            # Lưu cùng thư mục với input
            return str(input_path_obj.parent / md_filename)
        
        if not self.preserve_structure or base_dir is None:
            # Lưu flat vào output_dir
            return str(pathlib.Path(self.output_dir) / md_filename)
        
        # Giữ nguyên cấu trúc thư mục
        base_path = pathlib.Path(base_dir)
        try:
            relative_path = input_path_obj.parent.relative_to(base_path)
        except ValueError:
            # If not a child, use stem
            relative_path = "."
            
        output_subdir = pathlib.Path(self.output_dir) / relative_path
        output_subdir.mkdir(parents=True, exist_ok=True)
        
        return str(output_subdir / md_filename)
    
    def convert_file(self, file_path: str, output_path: Optional[str] = None, 
                    base_dir: Optional[str] = None) -> bool:
        """Convert một file sang Markdown"""
        file_path_obj = pathlib.Path(file_path)
        ext = file_path_obj.suffix.lower()
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            logger.warning(f"⚠️  Unsupported file type: {file_path}")
            return False
        
        # Xác định output path
        if output_path is None:
            output_path = self.get_output_path(file_path, base_dir)
        
        # Check if already exists
        if self.skip_existing and os.path.exists(output_path):
            logger.info(f"⏭️  Skipping (already exists): {output_path}")
            return True
        
        # Ensure output directory exists
        pathlib.Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # Convert based on extension
        if ext == '.pdf':
            success = self.convert_pdf_to_md(file_path, output_path)
        elif ext == '.docx':
            success = self.convert_docx_to_md(file_path, output_path)
        elif ext == '.doc':
            success = self.convert_doc_to_md(file_path, output_path)
        else:
            success = False
        
        return success

    def convert_folder(self, folder_path: str, recursive: bool = True, max_workers: Optional[int] = None) -> None:
        """Convert toàn bộ documents trong thư mục sử dụng multi-processing"""
        path = pathlib.Path(folder_path)
        
        if not path.exists():
            logger.error(f"❌ Folder không tồn tại: {folder_path}")
            return
        
        # Tìm tất cả files
        pattern = "**/*" if recursive else "*"
        all_files = []
        
        for ext in self.SUPPORTED_EXTENSIONS:
            files = list(path.glob(f"{pattern}{ext}"))
            all_files.extend(files)
        
        if not all_files:
            logger.warning(f"⚠️  Không tìm thấy file nào trong: {folder_path}")
            return
        
        # Filter out .doc files if .docx exists (ưu tiên .docx)
        filtered_files = []
        for file_info in all_files:
            if file_info.suffix.lower() == '.doc':
                docx_path = file_info.with_suffix('.docx')
                # Check if docx version exists in the same folder or in all_files
                if docx_path.exists() or docx_path in all_files:
                    logger.info(f"⏭️  Skipping {file_info.name} (using .docx version)")
                    continue
            filtered_files.append(str(file_info))
        
        num_files = len(filtered_files)
        if max_workers is None or max_workers <= 0:
            max_workers = min(multiprocessing.cpu_count(), num_files)
            
        logger.info(f"📂 Tìm thấy {num_files} file(s). Bắt đầu convert với {max_workers} workers...")
        logger.info("=" * 60)
        
        # Reset stats
        self.stats = {'total': 0, 'success': 0, 'failed': 0, 'skipped': 0}
        
        with ProcessPoolExecutor(max_workers=max_workers) as executor:
            futures = [
                executor.submit(
                    _worker_convert, 
                    f, self.output_dir, self.preserve_structure, 
                    self.skip_existing, folder_path
                ) for f in filtered_files
            ]
            
            for i, future in enumerate(as_completed(futures), 1):
                try:
                    res = future.result()
                    status = res['status']
                    fname = os.path.basename(res['file'])
                    
                    if status == 'skipped':
                        self.stats['skipped'] += 1
                    elif status == 'success':
                        self.stats['success'] += 1
                        self.stats['total'] += 1
                    else:
                        self.stats['failed'] += 1
                        self.stats['total'] += 1
                    
                    logger.info(f"[{i}/{num_files}] {status.upper()}: {fname}")
                except Exception as e:
                    logger.error(f"Error in worker: {str(e)}")
                    self.stats['failed'] += 1
                    self.stats['total'] += 1
        
        # Print summary
        logger.info("\n" + "=" * 60)
        logger.info("📊 CONVERSION SUMMARY")
        logger.info("=" * 60)
        logger.info(f"Total files attempted: {self.stats['total']}")
        logger.info(f"✅ Success: {self.stats['success']}")
        logger.info(f"❌ Failed: {self.stats['failed']}")
        logger.info(f"⏭️  Skipped: {self.stats['skipped']}")
        logger.info("=" * 60)

def _worker_convert(file_path: str, output_dir: Optional[str], preserve_structure: bool, 
                   skip_existing: bool, base_dir: Optional[str]) -> Dict:
    """Helper function for parallel processing (phải ở cấp module)"""
    converter = DocumentConverter(
        output_dir=output_dir,
        preserve_structure=preserve_structure,
        skip_existing=skip_existing
    )
    
    # Calculate output path to check skip early
    output_path = converter.get_output_path(file_path, base_dir)
    
    if skip_existing and os.path.exists(output_path):
        return {"status": "skipped", "file": file_path}
    
    success = converter.convert_file(file_path, output_path=output_path, base_dir=base_dir)
    return {"status": "success" if success else "failed", "file": file_path}

def main():
    """Main function với command-line arguments"""
    parser = argparse.ArgumentParser(
        description='Convert PDF/DOC/DOCX files to Markdown',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Convert một file
  python doc2md_enhanced.py input.pdf
  
  # Convert toàn bộ folder với multi-workers
  python doc2md_enhanced.py /path/to/folder --folder --workers 4
        """
    )
    
    parser.add_argument('input', help='Input file hoặc folder path')
    parser.add_argument('-o', '--output', help='Output file path (chỉ dùng cho single file)')
    parser.add_argument('--folder', action='store_true', help='Convert toàn bộ folder')
    parser.add_argument('--output-dir', help='Output directory cho folder conversion')
    parser.add_argument('--no-preserve', action='store_true', 
                       help='Không giữ cấu trúc thư mục (lưu flat)')
    parser.add_argument('--skip-existing', action='store_true', 
                       help='Bỏ qua file đã được convert')
    parser.add_argument('--no-recursive', action='store_true', 
                       help='Không tìm kiếm recursive trong subfolder')
    parser.add_argument('-w', '--workers', type=int, default=multiprocessing.cpu_count(),
                       help='Số lượng process chạy song song (mặc định: số CPU)')
    
    args = parser.parse_args()
    
    # Validate input
    if not os.path.exists(args.input):
        logger.error(f"❌ Input không tồn tại: {args.input}")
        sys.exit(1)
    
    # Create converter
    converter = DocumentConverter(
        output_dir=args.output_dir,
        preserve_structure=not args.no_preserve,
        skip_existing=args.skip_existing
    )
    
    # Convert
    if args.folder or os.path.isdir(args.input):
        converter.convert_folder(args.input, recursive=not args.no_recursive, max_workers=args.workers)
    else:
        converter.convert_file(args.input, output_path=args.output)


if __name__ == "__main__":
    main()
